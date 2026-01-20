#!/usr/bin/env python3
import sys
import os
import argparse
from docx import Document

def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading file: {str(e)}"

def create_docx(content_file, output_file):
    try:
        doc = Document()
        
        with open(content_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
                
        doc.save(output_file)
        return f"Document saved to {output_file}"
    except Exception as e:
        return f"Error creating file: {str(e)}"

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Read or create DOCX files")
    subparsers = parser.add_subparsers(dest="command", help="Command to execute")
    
    # Read command
    read_parser = subparsers.add_parser("read", help="Read text from a .docx file")
    read_parser.add_argument("file", help="Path to the .docx file")
    
    # Create command
    create_parser = subparsers.add_parser("create", help="Create a .docx file from markdown text")
    create_parser.add_argument("input", help="Input text/markdown file")
    create_parser.add_argument("output", help="Output .docx file path")
    
    args = parser.parse_args()
    
    if args.command == "read":
        if not os.path.exists(args.file):
            print(f"Error: File '{args.file}' not found.")
            sys.exit(1)
        print(read_docx(args.file))
        
    elif args.command == "create":
        if not os.path.exists(args.input):
            print(f"Error: Input file '{args.input}' not found.")
            sys.exit(1)
        print(create_docx(args.input, args.output))
        
    else:
        parser.print_help()
